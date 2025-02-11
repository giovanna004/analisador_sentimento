from docx import Document
from analisador_simples import *

def ler_arquivo(doc):
    return [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

   
def resumo_comentarios(status_lista):
    doc = Document()

    qtd_positivos = status_lista.count("Positivo.")
    qtd_negativos = status_lista.count("Negativo.")
    qtd_neutros = status_lista.count("Neutro.")
    total = qtd_positivos + qtd_negativos + qtd_neutros
       
    doc.add_heading('Relatório de Análise de Sentimentos', level=1)
    doc.add_paragraph(f'Total de comentários: {total}')
    
    if total > 0:
        doc.add_paragraph(f'Comentários Positivos: {(qtd_positivos/total) * 100:.2f}%')
        doc.add_paragraph(f'Comentários Negativos: {(qtd_negativos/total) * 100:.2f}%')
        doc.add_paragraph(f'Comentários Neutros: {(qtd_neutros/total) * 100:.2f}%')
    
    doc.save('resumo_avaliacoes.docx')


def main():
    doc = Document('avaliacoes.docx')

    positivas = ["bom", "gostei", "gostar", "gosto", "ótimo", "legal", "agrádavel", "alegre", "feliz", "excelente", "boa"]
    negativas = ["ruim", "péssimo", "horroroso", "chato", "odiei", "odiar", "triste", "desagradável", "horrível", "horror", "decepção", "decepcionado"]
    
    comentarios = ler_arquivo(doc)
    status_lista = []

    for comentario in comentarios:
        tokens = pre_processamento(comentario)
        status = classificar_sentimento(tokens, positivas, negativas)
        status_lista.append(status)
        print(f"Comentário: {comentario}\nClassificação: {status}\n")

    resumo_comentarios(status_lista)

main()
